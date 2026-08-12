"""Auxiliary capabilities for the Qwanto gateway: web search, document
extraction, voice transcription, and the Agent tool runtime.

All helpers are dependency-free at runtime. Optional support is only used when
the corresponding third-party package is importable.
"""

from __future__ import annotations

import hashlib
import io
import json
import os
import re
import struct
import subprocess
import sys
import time
import unicodedata
import urllib.parse
import urllib.request
from concurrent.futures import ThreadPoolExecutor
from typing import Any, Callable, Dict, List, Optional, Tuple

# -------- public limits (kept conservative on purpose) --------

MAX_SEARCH_RESULTS = 6
MAX_FETCH_BYTES = 256 * 1024
MAX_FILE_BYTES = 2 * 1024 * 1024
MAX_PDF_PAGES = 40
MAX_VOICE_BYTES = 24 * 1024 * 1024
AGENT_STEP_LIMIT = 8
SEARCH_TIMEOUT_SEC = 10.0
FETCH_TIMEOUT_SEC = 12.0

# -------- shared in-memory cache (process lifetime) --------

_CACHE: Dict[str, Tuple[float, Any]] = {}


def _cache_get(key: str, ttl: int) -> Optional[Any]:
    entry = _CACHE.get(key)
    if not entry:
        return None
    expiry, value = entry
    if time.time() > expiry:
        _CACHE.pop(key, None)
        return None
    return value


def _cache_put(key: str, ttl: int, value: Any) -> None:
    _CACHE[key] = (time.time() + ttl, value)


# ============================================================ web search

_USER_AGENT = "Qwanto/1.0 (+https://github.com/SaifHu98/Qwanto)"


def _http_get(url: str, *, timeout: float, headers: Optional[Dict[str, str]] = None) -> bytes:
    request = urllib.request.Request(url, headers={"User-Agent": _USER_AGENT, **(headers or {})})
    with urllib.request.urlopen(request, timeout=timeout) as response:
        return response.read()


def _strip_html(text: str) -> str:
    text = re.sub(r"(?is)<script\b.*?</script>", " ", text)
    text = re.sub(r"(?is)<style\b.*?</style>", " ", text)
    text = re.sub(r"(?is)<noscript\b.*?</noscript>", " ", text)
    text = re.sub(r"(?is)<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.replace("&nbsp;", " ").replace("&amp;", "&").replace("&quot;", '"').strip()


def _parse_duckduckgo(html: str) -> List[Dict[str, str]]:
    """Best-effort HTML parsing of DuckDuckGo's lite/HTML result page. The
    structure is not stable; we extract titles, snippets, and URLs conservatively.
    """
    results: List[Dict[str, str]] = []
    blocks = re.split(r"(?is)<!--\s*result\s+-->", html)
    for block in blocks:
        title_match = re.search(r'(?is)<a[^>]+class="result__a"[^>]*href="([^"]+)"[^>]*>(.*?)</a>', block)
        if not title_match:
            continue
        href = title_match.group(1)
        if href.startswith("//"):
            href = "https:" + href
        title = _strip_html(title_match.group(2))
        if not title:
            continue
        snippet_match = re.search(r'(?is)<a[^>]+class="result__snippet"[^>]*>(.*?)</a>', block)
        snippet = _strip_html(snippet_match.group(1)) if snippet_match else ""
        results.append({"title": title, "url": href, "snippet": snippet})
    if not results:
        # Fallback: scrape anchors + nearest text node.
        for match in re.finditer(r'(?is)<a[^>]+href="(https?://[^"]+)"[^>]*>(.*?)</a>', html):
            href = match.group(1)
            if any(domain in href for domain in ("duckduckgo.com", "duck.com")):
                continue
            title = _strip_html(match.group(2))
            if len(title) < 8:
                continue
            results.append({"title": title[:160], "url": href, "snippet": ""})
            if len(results) >= MAX_SEARCH_RESULTS:
                break
    return results[:MAX_SEARCH_RESULTS]


def web_search(query: str, *, max_results: int = MAX_SEARCH_RESULTS) -> List[Dict[str, str]]:
    """Search DuckDuckGo's HTML endpoint. Returns at most max_results items.

    Each item: {"title", "url", "snippet"}.
    """
    if not query or not query.strip():
        return []
    query = query.strip()[:256]
    key = "search:" + hashlib.sha256(query.lower().encode("utf-8")).hexdigest()[:16]
    cached = _cache_get(key, ttl=600)
    if cached is not None:
        return cached[:max_results]
    encoded = urllib.parse.quote_plus(query)
    url = f"https://html.duckduckgo.com/html/?q={encoded}&kl=us-en"
    try:
        html = _http_get(url, timeout=SEARCH_TIMEOUT_SEC,
                         headers={"Accept": "text/html"})
    except Exception:
        return []
    text = html.decode("utf-8", "replace")
    parsed = _parse_duckduckgo(text)
    _cache_put(key, ttl=600, value=parsed)
    return parsed[:max_results]


def web_fetch(url: str, *, max_bytes: int = MAX_FETCH_BYTES) -> Dict[str, Any]:
    """Fetch a URL and return its title + cleaned text. HTML responses are
    stripped; JSON responses are surfaced as-is. Other content types return a
    short metadata blob and a body excerpt.
    """
    if not url or not url.startswith(("http://", "https://")):
        return {"url": url, "error": "unsupported_url"}
    key = "fetch:" + hashlib.sha256(url.encode("utf-8")).hexdigest()[:16]
    cached = _cache_get(key, ttl=300)
    if cached is not None:
        return cached
    try:
        raw = _http_get(url, timeout=FETCH_TIMEOUT_SEC,
                        headers={"Accept": "text/html,application/json;q=0.9,*/*;q=0.5"})
    except Exception as exc:
        return {"url": url, "error": f"fetch_failed: {exc.__class__.__name__}"}
    body = raw[:max_bytes]
    content_type = ""
    try:
        with urllib.request.urlopen(urllib.request.Request(url, headers={"User-Agent": _USER_AGENT}),
                                     timeout=FETCH_TIMEOUT_SEC) as response:
            content_type = response.headers.get("Content-Type", "")
    except Exception:
        pass
    result: Dict[str, Any] = {"url": url, "content_type": content_type, "bytes": len(body)}
    lowered = content_type.lower()
    if "json" in lowered:
        try:
            payload = json.loads(body.decode("utf-8", "replace"))
        except Exception as exc:
            result["error"] = f"json_parse_failed: {exc.__class__.__name__}"
            return result
        if isinstance(payload, (dict, list)):
            result["json"] = payload
        else:
            result["text"] = str(payload)
    elif "html" in lowered or b"<" in body[:2048]:
        text = body.decode("utf-8", "replace")
        title_match = re.search(r"(?is)<title[^>]*>(.*?)</title>", text)
        title = _strip_html(title_match.group(1)) if title_match else ""
        clean = _strip_html(text)
        result["title"] = title[:200]
        result["text"] = clean[:max_bytes]
    else:
        try:
            result["text"] = body.decode("utf-8", "replace")
        except UnicodeDecodeError:
            result["text"] = ""
            result["note"] = "binary content; not transcribed"
    _cache_put(key, ttl=300, value=result)
    return result


# ============================================================ document extraction


def _decode_bytes(data: bytes) -> str:
    for codec in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(codec)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", "replace")


def _normalize_text(text: str, max_chars: int) -> str:
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    if len(text) > max_chars:
        text = text[:max_chars] + "\n…[truncated]"
    return text.strip()


def _read_text_file(data: bytes) -> str:
    return _normalize_text(_decode_bytes(data), max_chars=120_000)


def _try_import_pypdf():
    try:
        import pypdf  # type: ignore
        return pypdf
    except Exception:
        return None


def _try_import_docx():
    try:
        import docx  # type: ignore
        return docx
    except Exception:
        return None


def _extract_pdf(data: bytes) -> str:
    pypdf = _try_import_pypdf()
    if pypdf is not None:
        try:
            reader = pypdf.PdfReader(io.BytesIO(data))
            parts: List[str] = []
            for index, page in enumerate(reader.pages):
                if index >= MAX_PDF_PAGES:
                    parts.append(f"\n[stopped at page {MAX_PDF_PAGES}]")
                    break
                try:
                    text = page.extract_text() or ""
                except Exception:
                    text = ""
                if text:
                    parts.append(text)
            return _normalize_text("\n\n".join(parts), max_chars=120_000)
        except Exception:
            pass
    # Fallback: a very small embedded PDF reader that handles uncompressed
    # text streams. PDF is complex; without pypdf we surface a hint instead of
    # producing garbled text.
    return ("[PDF text extraction requires the optional 'pypdf' package. "
            "Install it with: pip install pypdf]")


def _extract_docx(data: bytes) -> str:
    docx = _try_import_docx()
    if docx is None:
        return ("[DOCX text extraction requires the optional 'python-docx' "
                "package. Install it with: pip install python-docx]")
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        return f"[DOCX open failed: {exc}]"
    parts: List[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells if cell.text]
            if cells:
                parts.append(" | ".join(cells))
    return _normalize_text("\n".join(parts), max_chars=120_000)


def extract_attachment(name: str, data: bytes, mime: str) -> Dict[str, Any]:
    """Extract text from an uploaded attachment. Returns {name, kind, text, note}.
    `kind` is one of: "text", "code", "pdf", "docx", "image", "binary", "empty".
    """
    info: Dict[str, Any] = {"name": name, "bytes": len(data), "mime": mime or ""}
    if not data:
        info["kind"] = "empty"
        info["text"] = ""
        return info
    if len(data) > MAX_FILE_BYTES:
        info["kind"] = "binary"
        info["note"] = f"file exceeds {MAX_FILE_BYTES} bytes; not transcribed"
        info["text"] = ""
        return info
    lowered = (mime or "").lower()
    by_extension = name.lower()
    if lowered.startswith("image/") or by_extension.endswith((
            ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".svg")):
        info["kind"] = "image"
        info["note"] = ("images are not parsed; the model receives the filename "
                        "and MIME type. Run an OCR pass locally if you need text.")
        info["text"] = ""
        return info
    if lowered == "application/pdf" or by_extension.endswith(".pdf"):
        info["kind"] = "pdf"
        info["text"] = _extract_pdf(data)
        return info
    if lowered in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                   "application/msword") or by_extension.endswith((".docx", ".doc")):
        info["kind"] = "docx"
        info["text"] = _extract_docx(data)
        return info
    if lowered.startswith("text/") or by_extension.endswith((
            ".txt", ".md", ".csv", ".tsv", ".json", ".jsonl", ".yaml", ".yml",
            ".xml", ".html", ".htm", ".log", ".ini", ".toml", ".sql", ".py",
            ".js", ".ts", ".tsx", ".jsx", ".c", ".cc", ".cpp", ".cxx", ".h",
            ".hpp", ".rs", ".go", ".java", ".kt", ".rb", ".php", ".sh", ".ps1",
            ".bat", ".css", ".scss", ".sass")):
        info["kind"] = "code" if by_extension.endswith((
            ".py", ".js", ".ts", ".tsx", ".jsx", ".c", ".cc", ".cpp", ".cxx",
            ".h", ".hpp", ".rs", ".go", ".java", ".kt", ".rb", ".php", ".sh",
            ".ps1", ".bat")) else "text"
        info["text"] = _read_text_file(data)
        return info
    # Unknown binary: best-effort utf-8 peek.
    try:
        preview = _decode_bytes(data[:2048])
    except Exception:
        info["kind"] = "binary"
        info["text"] = ""
        return info
    if "\0" in preview:
        info["kind"] = "binary"
        info["text"] = ""
        return info
    info["kind"] = "text"
    info["text"] = _read_text_file(data)
    return info


# ============================================================ voice transcription


def _try_import_whisper_cpp():
    try:
        import whisper_cpp  # type: ignore
        return whisper_cpp
    except Exception:
        return None


def _whisper_cli() -> Optional[str]:
    for candidate in ("whisper", "whisper-cpp", "main"):
        path = _which(candidate)
        if path:
            return path
    return None


def _which(name: str) -> Optional[str]:
    from shutil import which
    return which(name)


def transcribe_audio(name: str, data: bytes, mime: str) -> Dict[str, Any]:
    """Transcribe a voice clip. Tries, in order:
      1. whisper-cpp Python bindings
      2. whisper.cpp CLI (`whisper` / `whisper-cpp` / `main`)
      3. built-in stub: returns a note so the UI can prompt the user.

    Returns: {name, mime, bytes, text, engine}.
    """
    if not data:
        return {"name": name, "mime": mime, "bytes": 0, "text": "",
                "engine": "none", "note": "empty audio payload"}
    if len(data) > MAX_VOICE_BYTES:
        return {"name": name, "mime": mime, "bytes": len(data), "text": "",
                "engine": "none", "note": f"audio exceeds {MAX_VOICE_BYTES} bytes"}
    whisper = _try_import_whisper_cpp()
    if whisper is not None and hasattr(whisper, "transcribe"):
        try:
            with tempfile_for_bytes(data, suffix=_suffix_for(name, mime)) as path:
                result = whisper.transcribe(path)
                text = result.get("text", "") if isinstance(result, dict) else str(result)
                return {"name": name, "mime": mime, "bytes": len(data),
                        "text": text.strip(), "engine": "whisper_cpp"}
        except Exception:
            pass
    cli = _whisper_cli()
    if cli:
        try:
            with tempfile_for_bytes(data, suffix=_suffix_for(name, mime)) as path:
                completed = subprocess.run(
                    [cli, "--model", os.environ.get("QWANTO_WHISPER_MODEL", "base"),
                     "--output-txt", "--no-timestamps", path],
                    capture_output=True, text=True, timeout=120)
            text_path = os.path.splitext(path)[0] + ".txt"
            if os.path.exists(text_path):
                with open(text_path, "r", encoding="utf-8", errors="replace") as handle:
                    text = handle.read()
                return {"name": name, "mime": mime, "bytes": len(data),
                        "text": text.strip(), "engine": os.path.basename(cli)}
        except Exception:
            pass
    return {"name": name, "mime": mime, "bytes": len(data), "text": "",
            "engine": "none",
            "note": ("no local whisper runtime found; install whisper-cpp Python "
                     "bindings, the whisper.cpp CLI on PATH, or set QWANTO_WHISPER_MODEL "
                     "and supply a compatible engine. Browser-side transcription is also "
                     "available as a fallback.").strip()}


import contextlib
import tempfile


@contextlib.contextmanager
def tempfile_for_bytes(data: bytes, suffix: str):
    handle = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    try:
        handle.write(data)
        handle.close()
        yield handle.name
    finally:
        try:
            os.unlink(handle.name)
        except OSError:
            pass


def _suffix_for(name: str, mime: str) -> str:
    lowered = (mime or "").lower()
    for ext in (".wav", ".mp3", ".m4a", ".ogg", ".flac", ".webm", ".mp4"):
        if lowered.endswith(ext) or name.lower().endswith(ext):
            return ext
    if name.lower().endswith((".wav", ".mp3", ".m4a", ".ogg", ".flac", ".webm")):
        return "." + name.rsplit(".", 1)[-1]
    return ".wav"


# ============================================================ agent tools

AGENT_TOOLS: Dict[str, Dict[str, Any]] = {
    "web_search": {
        "description": "Search the public web (DuckDuckGo HTML). Returns titles, URLs, and snippets.",
        "schema": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "The search query."},
            },
            "required": ["query"],
        },
    },
    "web_fetch": {
        "description": "Fetch a URL and return its cleaned text or JSON payload.",
        "schema": {
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "The URL to fetch."},
            },
            "required": ["url"],
        },
    },
    "datetime": {
        "description": "Return the current date and time in UTC and local time.",
        "schema": {"type": "object", "properties": {}},
    },
    "calculator": {
        "description": "Evaluate a small arithmetic expression (integer or decimal).",
        "schema": {
            "type": "object",
            "properties": {
                "expression": {"type": "string", "description": "Arithmetic expression."},
            },
            "required": ["expression"],
        },
    },
    "unit_convert": {
        "description": "Convert a numeric value between common units (length, mass, temperature, data).",
        "schema": {
            "type": "object",
            "properties": {
                "value": {"type": "number", "description": "Numeric value to convert."},
                "from_unit": {"type": "string"},
                "to_unit": {"type": "string"},
            },
            "required": ["value", "from_unit", "to_unit"],
        },
    },
    "plan": {
        "description": "Persist a structured plan (objective, steps, status) for a multi-step task.",
        "schema": {
            "type": "object",
            "properties": {
                "objective": {"type": "string"},
                "steps": {"type": "array", "items": {"type": "string"}},
            },
            "required": ["objective", "steps"],
        },
    },
}


def agent_tool_schemas() -> List[Dict[str, Any]]:
    return [{"type": "function",
             "function": {"name": name, "description": info["description"],
                          "parameters": info["schema"]}}
            for name, info in AGENT_TOOLS.items()]


def _calculator(expression: str) -> Dict[str, Any]:
    if not expression or not expression.strip():
        return {"ok": False, "error": "empty expression"}
    cleaned = expression.strip()
    if len(cleaned) > 200:
        return {"ok": False, "error": "expression too long"}
    if not re.fullmatch(r"[0-9\s\+\-\*\/\(\)\.\,eE]+", cleaned):
        return {"ok": False, "error": "unsupported characters"}
    try:
        value = eval(cleaned, {"__builtins__": {}}, {})  # noqa: S307 - sandboxed input
    except Exception as exc:
        return {"ok": False, "error": f"evaluation failed: {exc}"}
    return {"ok": True, "value": value}


def _unit_convert(value: float, from_unit: str, to_unit: str) -> Dict[str, Any]:
    factors = {
        "length": {"mm": 0.001, "cm": 0.01, "m": 1.0, "km": 1000.0,
                   "inch": 0.0254, "ft": 0.3048, "yard": 0.9144, "mile": 1609.344},
        "mass": {"g": 0.001, "kg": 1.0, "lb": 0.45359237, "oz": 0.0283495231,
                 "ton": 1000.0},
        "data": {"B": 1, "KB": 1024, "MB": 1024 ** 2, "GB": 1024 ** 3,
                 "TB": 1024 ** 4, "KiB": 1024, "MiB": 1024 ** 2,
                 "GiB": 1024 ** 3, "TiB": 1024 ** 4},
    }
    temperature = {"C": (0.0, -273.15), "F": (1.0, -459.67), "K": (0.0, 0.0)}
    try:
        if from_unit in temperature and to_unit in temperature:
            celsius = (value - 32) * 5 / 9 if from_unit == "F" else (value - 273.15) if from_unit == "K" else value
            if to_unit == "C":
                return {"ok": True, "value": celsius}
            if to_unit == "F":
                return {"ok": True, "value": celsius * 9 / 5 + 32}
            return {"ok": True, "value": celsius + 273.15}
        for category, table in factors.items():
            if from_unit in table and to_unit in table:
                base = value * table[from_unit]
                return {"ok": True, "value": base / table[to_unit]}
        return {"ok": False, "error": f"unsupported units: {from_unit} -> {to_unit}"}
    except Exception as exc:
        return {"ok": False, "error": str(exc)}


def call_agent_tool(name: str, arguments: Dict[str, Any]) -> Dict[str, Any]:
    """Dispatch an agent tool call. Returns a JSON-serializable dict."""
    if name not in AGENT_TOOLS:
        return {"ok": False, "error": f"unknown tool: {name}"}
    try:
        if name == "web_search":
            return {"ok": True, "results": web_search(arguments.get("query", ""))}
        if name == "web_fetch":
            return {"ok": True, "result": web_fetch(arguments.get("url", ""))}
        if name == "datetime":
            now = time.time()
            return {"ok": True, "utc": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime(now)),
                    "epoch": int(now)}
        if name == "calculator":
            return _calculator(arguments.get("expression", ""))
        if name == "unit_convert":
            return _unit_convert(float(arguments.get("value", 0)),
                                 str(arguments.get("from_unit", "")),
                                 str(arguments.get("to_unit", "")))
        if name == "plan":
            return {"ok": True, "saved": True,
                    "objective": arguments.get("objective", ""),
                    "steps": list(arguments.get("steps") or [])}
    except Exception as exc:
        return {"ok": False, "error": f"{exc.__class__.__name__}: {exc}"}
    return {"ok": False, "error": "unhandled tool"}


# ============================================================ plan/agent loop


def run_agent(user_request: str,
              history: List[Dict[str, Any]],
              executor: Callable[[Dict[str, Any]], Dict[str, Any]]) -> Dict[str, Any]:
    """Drive a multi-step agent loop.

    The executor is expected to call the underlying model and return a dict
    matching the shape `{"type": "text|tool", "content": ..., "tool": ..., "args": ...}`.

    The loop:
      1. Sends the user request + tool catalog to the executor.
      2. If the model returns a tool call, executes it through `call_agent_tool`,
         appends the result, and re-invokes the executor. Repeats up to
         AGENT_STEP_LIMIT times or until a final text answer is produced.
      3. Returns the transcript with the final text and any tool observations.
    """
    transcript: List[Dict[str, Any]] = []
    messages: List[Dict[str, Any]] = list(history or [])
    messages.append({"role": "user", "content": user_request})
    plan: List[str] = []
    final_text: str = ""
    steps = 0
    while steps < AGENT_STEP_LIMIT:
        steps += 1
        response = executor({
            "messages": messages,
            "tools": agent_tool_schemas(),
            "plan": plan,
            "step": steps,
            "limit": AGENT_STEP_LIMIT,
        }) or {}
        kind = response.get("type")
        if kind == "text":
            final_text = response.get("content", "") or final_text
            transcript.append({"role": "assistant", "type": "text", "content": final_text})
            break
        if kind == "tool":
            tool = response.get("tool") or ""
            args = response.get("args") or {}
            transcript.append({"role": "assistant", "type": "tool", "tool": tool, "args": args,
                               "step": steps})
            observation = call_agent_tool(tool, args)
            transcript.append({"role": "tool", "name": tool, "result": observation, "step": steps})
            messages.append({"role": "assistant", "type": "tool", "tool": tool, "args": args})
            messages.append({"role": "tool", "name": tool, "content": observation})
            if tool == "plan":
                plan = list(args.get("steps") or [])
            continue
        final_text = response.get("content", "") if response else ""
        if not final_text:
            final_text = "The agent did not return a final response."
        transcript.append({"role": "assistant", "type": "text", "content": final_text})
        break
    if not final_text:
        final_text = "The agent exceeded the step limit without a final response."
    return {"final_text": final_text, "transcript": transcript, "plan": plan,
            "steps_used": steps}


# ============================================================ render helpers

def render_research(results: List[Dict[str, str]]) -> str:
    """Render a search result list as a model-readable text block."""
    if not results:
        return "(no search results)"
    parts: List[str] = []
    for index, item in enumerate(results, start=1):
        title = item.get("title", "(untitled)").strip()
        url = item.get("url", "").strip()
        snippet = item.get("snippet", "").strip()
        parts.append(f"[{index}] {title}\n    {url}\n    {snippet}")
    return "\n".join(parts)


def build_attachments_block(attachments: List[Dict[str, Any]]) -> str:
    """Render the attachment list as a model-readable text block."""
    if not attachments:
        return ""
    parts = ["[Attached files]"]
    for item in attachments:
        name = item.get("name", "(unknown)")
        kind = item.get("kind", "binary")
        note = item.get("note", "")
        text = item.get("text", "")
        header = f"--- {name} ({kind})"
        if note:
            header += f" — {note}"
        if text:
            parts.append(f"{header}\n{text}")
        else:
            parts.append(f"{header}\n(no extracted text)")
    return "\n\n".join(parts)
